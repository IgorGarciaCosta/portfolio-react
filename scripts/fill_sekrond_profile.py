"""
Fills the SEKROND consultant-profile template from the Work Experience section
onward, using Igor's real data (from SEKROND_CV_Fill_Guide.md), while preserving
the template's exact formatting.

Strategy:
- Edit existing placeholder paragraphs in place (set run text, keep run/paragraph
  formatting) so styling is untouched.
- Where more bullet lines are needed than placeholders exist, clone the placeholder
  paragraph's XML (deep copy) so the new lines inherit identical formatting.
- Rebuild the "Software/Tech stack" table by cloning the template's own header /
  data / blank rows as exemplars, so every added row matches the original design.

Leaves everything BEFORE Work Experience (Name, Title, Profile, Key Skills) intact.
"""
import copy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\CV_Sekrond_Template_S.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\IgorGarcia_SEKROND_Profile.docx"

doc = Document(INPUT)
paras = doc.paragraphs  # snapshot; held Paragraph refs stay valid across inserts


# ---------- helpers ----------
def set_all(para, text):
    """Put text in the first run, clear the rest (keeps first run's formatting)."""
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def clear_all(para):
    for r in para.runs:
        r.text = ""


def set_run(para, idx, text):
    para.runs[idx].text = text


def append_value(para, text):
    """Append a value run (inherits paragraph/style formatting -> non-bold value)."""
    para.add_run(text)


def clone_after(ref_para, text=None):
    """Deep-copy ref paragraph, insert right after it, optionally set its text."""
    new_el = copy.deepcopy(ref_para._p)
    # a cloned body paragraph must never inherit a section break
    pPr = new_el.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            pPr.remove(sectPr)
    ref_para._p.addnext(new_el)
    p = Paragraph(new_el, ref_para._parent)
    if text is not None:
        set_all(p, text)
    return p


# ---------- anchor guards (fail fast if the template shifted) ----------
assert "Ford Motor Company" in paras[46].text, paras[46].text
assert paras[61].text.strip().startswith("Company: Company Name"), paras[61].text
assert "Company Name" in paras[73].text, paras[73].text
assert paras[89].text.strip().startswith("Company"), paras[89].text
assert paras[105].text.strip().startswith("Title"), paras[105].text
assert paras[120].text.strip().startswith("Program"), paras[120].text
assert paras[143].text.strip().endswith("Advanced"), paras[143].text


# =====================================================================
# 5.1  FORD — Software Engineer / C++/Python Developer (block 1, already started)
# =====================================================================
# Fix the stray "Text" run inside the hand-tracking achievement (P51 run idx 2)
if paras[51].runs[2].text == "Text":
    paras[51].runs[2].text = ""

# Add the two strongest metric achievements after the existing occlusion one (P52)
a1 = clone_after(
    paras[52],
    "Engineered high-performance internal tooling and automated validation pipelines "
    "in C++, replacing physical prototypes and contributing to an estimated ~40% "
    "annual material-cost reduction.")
clone_after(
    a1,
    "Cut design iteration cycles by ~50% by working directly with internal customers "
    "(Design and Engineering) to define and build the tools they needed.")

# Role Description bullets (P55, P56 = "Text") + one extra
set_all(paras[55],
        "Develop and optimize high-density real-time visualizations across multiple "
        "hardware targets, sustaining stable frame rates while processing 5M+ "
        "polygons / data points.")
set_all(paras[56],
        "Automate build, test and deployment with CI/CD (GitHub Actions) for VR "
        "validation tooling across 5+ platforms (HTC Vive, Meta Quest Pro).")
clone_after(paras[56],
            "Focus on performance profiling, memory management, stability and "
            "long-term maintainability.")

# Software / Tools (P58)
append_value(paras[58],
             " C++, Unreal Engine 5 (C++/Blueprint), Python, GitHub Actions, "
             "Git/Perforce, computer vision, performance profilers, HTC Vive, "
             "Meta Quest Pro, CATIA data.")


# =====================================================================
# 5.2  FORD — Virtual Reality Researcher (block 2)
# =====================================================================
set_run(paras[61], 1, "Ford Motor Company")
set_run(paras[62], 3, "Virtual Reality Researcher")
set_run(paras[63], 1, "08/2021 \u2013 12/2022")
set_all(paras[66],
        "Spearheaded R&D for automotive VR applications, establishing workflows "
        "later adopted as the company standard for virtual prototyping.")
set_all(paras[67],
        "Improved user-assessment accuracy by ~20% by translating physical "
        "requirements into virtual specifications with Design and Engineering teams.")
set_all(paras[70],
        "Built initial prototypes for virtual user-experience testing, validating "
        "key interaction concepts for future production use.")
set_all(paras[71],
        "Collaborated with Design and Engineering teams to turn physical "
        "requirements into testable virtual specifications.")
append_value(paras[72],
             " Unreal Engine, VR hardware (HTC Vive), C++, 3D pipeline tools.")


# =====================================================================
# 5.3  CAFUNDÓ — Software Engineer (C# / AI Integration) (block 3)
# =====================================================================
set_run(paras[73], 3, "Cafund\u00f3 Creative Studio  (Contract)")
set_run(paras[74], 3, "Software Engineer (C# / AI Integration)")
set_run(paras[75], 1, "06/2025 \u2013 Present")
set_all(paras[78],
        "Architected an AI-powered interactive totem (\u201cSelfie with Lis\u201d) in "
        "C#/Unity for high-traffic public spaces, serving 500+ daily interactions "
        "with zero downtime.")
set_all(paras[79],
        "Reduced voice-to-response latency by 30% by integrating OpenAI and Azure "
        "Cognitive Services into a resilient real-time inference/response pipeline "
        "(auth, rate-limit and failure handling).")
set_all(paras[82],
        "Designed real-time, event-driven interaction systems for seamless public use.")
set_all(paras[83],
        "Built accessible UI/UX and handled auth, rate limits and failure scenarios "
        "for reliable, always-on operation.")
append_value(paras[84],
             " C#, Unity, OpenAI API, Azure Cognitive Services, REST APIs, "
             "real-time inference.")


# =====================================================================
# 6.  INTERNATIONAL EXPERIENCE — Blue Gravity Studios
# =====================================================================
set_run(paras[89], 2, "Blue Gravity Studios  (Contract)")
set_run(paras[90], 2, "Software Engineer (C++ / Unreal Engine 5)")
set_run(paras[91], 2, "01/2024 \u2013 09/2024")
set_all(paras[94],
        "Engineered core network replication logic (TCP/UDP) for a multiplayer "
        "architecture in C++, optimizing bandwidth usage by ~25% to support highly "
        "concurrent online sessions.")
set_all(paras[95],
        "Shipped the title \u201cSkateNation XL\u201d to PlayStation 5 and Xbox "
        "Series X|S.")
b1 = clone_after(
    paras[95],
    "Reduced bug-fixing overhead by 20% by improving code modularity across core "
    "software systems.")
clone_after(
    b1,
    "Cut input latency by 15ms and reached 100% cross-platform compatibility with a "
    "scalable keyboard + gamepad input system.")
set_all(paras[98],
        "Collaborated with a distributed, multidisciplinary international team to "
        "optimize draw calls and texture memory, improving performance on lower-end "
        "hardware.")
set_all(paras[99],
        "Worked in a fast-paced console production pipeline (PS5 & Xbox Series X|S) "
        "using Perforce/Git.")


# =====================================================================
# 8.  PUBLICATIONS  ->  three separate published tools (2 FAB plugins + Blender)
# =====================================================================
# Capture pristine copies of one publication block (Title / Publisher / Date /
# blank / Description) so each new entry inherits identical formatting.
pub_title_tmpl = copy.deepcopy(paras[105]._p)
pub_pubr_tmpl = copy.deepcopy(paras[106]._p)
pub_date_tmpl = copy.deepcopy(paras[107]._p)
pub_blank_tmpl = copy.deepcopy(paras[108]._p)
pub_desc_tmpl = copy.deepcopy(paras[109]._p)


def _fill_pub(title_p, pubr_p, date_p, desc_p, title, publisher, date_val, desc):
    # keep bold labels ("Title:", "Publisher", "Date:"), only replace the values
    set_run(title_p, 2, title)
    if len(title_p.runs) > 3:
        set_run(title_p, 3, "")
    set_run(pubr_p, 1, ": " + publisher + " (")
    set_run(pubr_p, 2, date_val)
    if len(pubr_p.runs) > 3:
        set_run(pubr_p, 3, ")")
    set_run(date_p, 2, date_val)
    set_all(desc_p, desc)


def _insert_after(ref_p, tmpl_el):
    new_el = copy.deepcopy(tmpl_el)
    ref_p._p.addnext(new_el)
    return Paragraph(new_el, ref_p._parent)


PUBS = [
    ("Professional Measurement Tool \u2014 Unreal Engine Editor Plugin",
     "FAB / Unreal Engine Marketplace", "2025 \u2013 2026",
     "Spline-based in-editor measurement plugin for Unreal Engine (point-to-point "
     "distances, enclosed areas and three snap modes), delivered as a custom Editor "
     "Mode. Commercially published on the FAB / Epic marketplace. "
     "Repo: github.com/IgorGarciaCosta."),
    ("Runtime Skinned Mesh Exporter \u2014 Unreal Engine Plugin",
     "FAB / Unreal Engine Marketplace", "2025 \u2013 2026",
     "Runtime skinned-mesh exporter built on a fully asynchronous GPU\u2192CPU "
     "readback pipeline, supporting UE 4.27\u20135.6. Commercially published on the "
     "FAB / Epic marketplace. Repo: github.com/IgorGarciaCosta."),
    ("Smart Mesh Cleaner Pro \u2014 Blender Add-on",
     "Open-source (GitHub)", "2025",
     "Blender add-on that streamlines asset maintenance with a Smart Trash Bin "
     "system \u2014 safely staging deletions, previewing impact and enabling "
     "one-click restore, so artists can clean scenes confidently without data loss. "
     "Repo: github.com/IgorGarciaCosta/blender-smart-mesh-cleaner."),
]

# Entry 1 reuses the existing template paragraphs in place.
_fill_pub(paras[105], paras[106], paras[107], paras[109], *PUBS[0])

# Entries 2..N: clone the whole block (with a blank separator) after the previous
# description paragraph, then fill it.
ref = paras[109]
for title, publisher, date_val, desc in PUBS[1:]:
    ref = _insert_after(ref, pub_blank_tmpl)          # separator blank line
    t_p = _insert_after(ref, pub_title_tmpl); ref = t_p
    pubr_p = _insert_after(ref, pub_pubr_tmpl); ref = pubr_p
    date_p = _insert_after(ref, pub_date_tmpl); ref = date_p
    ref = _insert_after(ref, pub_blank_tmpl)          # blank before description
    desc_p = _insert_after(ref, pub_desc_tmpl); ref = desc_p
    _fill_pub(t_p, pubr_p, date_p, desc_p, title, publisher, date_val, desc)


# =====================================================================
# 9.  SOFTWARE / TECH STACK  — rebuild table from its own row exemplars
# =====================================================================
tbl = doc.tables[0]
header_tr = copy.deepcopy(tbl.rows[10]._tr)   # a category header row
data_tr = copy.deepcopy(tbl.rows[11]._tr)     # a data row
blank_tr = copy.deepcopy(tbl.rows[6]._tr)     # a blank spacer row

# remove every existing row (keeps tblPr / tblGrid)
for row in list(tbl.rows):
    row._tr.getparent().remove(row._tr)


def _set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def add_header(title):
    tr = copy.deepcopy(header_tr)
    tbl._tbl.append(tr)
    row = tbl.rows[-1]
    _set_cell(row.cells[0], title)
    _set_cell(row.cells[1], "Experience (yrs)")
    _set_cell(row.cells[2], "Skill (1\u20135)")
    _set_cell(row.cells[3], "Last used")


def add_data(name, yrs, skill, last):
    tr = copy.deepcopy(data_tr)
    tbl._tbl.append(tr)
    row = tbl.rows[-1]
    _set_cell(row.cells[0], name)
    _set_cell(row.cells[1], yrs)
    _set_cell(row.cells[2], skill)
    _set_cell(row.cells[3], last)


def add_blank():
    tbl._tbl.append(copy.deepcopy(blank_tr))


W = "Within days"
STACK = [
    ("Concepts & Methodology", [
        ("Object-Oriented Programming", "5", "4", W),
        ("Software Development (SDLC)", "5", "4", W),
        ("Clean Architecture / SOLID", "3", "4", W),
        ("Agile / Scrum", "4", "3", W),
        ("Performance Profiling & Optimization", "5", "4", W),
        ("CI/CD & Automation", "3", "3", W),
        ("Design Patterns", "3", "4", W),
    ]),
    ("IDEs, Debuggers & Compilers", [
        ("Visual Studio", "5", "4", W),
        ("Visual Studio Code", "5", "4", W),
        ("Rider / JetBrains", "2", "3", W),
        ("Unreal Editor", "5", "4", W),
        ("Unity Editor", "2", "3", W),
    ]),
    ("Programming Languages", [
        ("C++", "5", "4", W),
        ("C#", "3", "4", W),
        ("Python", "3", "3", W),
        ("TypeScript", "3", "3", W),
        ("JavaScript", "3", "3", W),
        ("SQL", "3", "3", W),
        ("HTML / CSS", "3", "3", W),
    ]),
    ("Frameworks & Runtimes", [
        ("Unreal Engine 5 (C++/Blueprint)", "5", "4", W),
        (".NET / ASP.NET Core", "3", "4", W),
        ("Entity Framework Core", "2", "3", W),
        ("React", "3", "3", W),
        ("Node.js", "2", "3", W),
        ("Unity", "2", "3", W),
        ("Tailwind CSS", "2", "3", W),
    ]),
    ("AI / ML Integration", [
        ("OpenAI API", "2", "3", W),
        ("Azure Cognitive Services", "1", "3", W),
        ("Google Gemini", "1", "3", W),
        ("Model Context Protocol (MCP)", "1", "3", W),
        ("Computer Vision (MediaPipe, OpenCV)", "2", "3", W),
    ]),
    ("Cloud Infrastructure & Services", [
        ("AWS Lambda", "2", "3", W),
        ("AWS API Gateway", "2", "3", W),
        ("AWS SQS", "1", "3", W),
        ("AWS DynamoDB", "1", "3", W),
        ("Azure (Cognitive Services)", "1", "3", W),
        ("Oracle Cloud", "1", "2", "2026"),
        ("Firebase / Firestore", "1", "3", "2025"),
    ]),
    ("DevOps & Source Control", [
        ("Git", "5", "4", W),
        ("GitHub Actions (CI/CD)", "3", "3", W),
        ("Docker / Docker Compose", "2", "3", W),
        ("Perforce", "2", "3", "2024"),
        ("Linux fundamentals", "4", "3", W),
    ]),
    ("Databases", [
        ("PostgreSQL", "2", "3", W),
        ("SQLite", "2", "3", W),
        ("DynamoDB", "1", "3", "2025"),
        ("Firebase / Firestore", "1", "3", "2025"),
    ]),
    ("Standards & Protocols", [
        ("REST / HTTP", "3", "4", W),
        ("OpenAPI / Swagger", "2", "3", W),
        ("TCP / UDP (network replication)", "3", "3", W),
        ("JWT / cookie auth", "2", "3", W),
        ("ISO 26262 (awareness, automotive)", "2", "2", W),
    ]),
]

for i, (title, rows) in enumerate(STACK):
    add_header(title)
    for name, yrs, skill, last in rows:
        add_data(name, yrs, skill, last)
    if i != len(STACK) - 1:
        add_blank()


# =====================================================================
# 10.  EDUCATION
# =====================================================================
set_run(paras[120], 1, "Computer Engineering (Bachelor's degree)")
set_run(paras[121], 1, ": State University of Feira de Santana (UEFS)")
set_run(paras[122], 1, "Feira de Santana, Bahia")
set_run(paras[122], 3, "Brazil")
# remove the unused 2nd program placeholder + its description
clear_all(paras[124])
clear_all(paras[125])
clear_all(paras[126])
clear_all(paras[128])


# =====================================================================
# 11.  TRAININGS / COURSES  ->  completed online courses (Udemy + LinkedIn)
#      Format: "Course Name, Issuer. (Year)"
# =====================================================================
COURSES = [
    "Designing for Virtual Reality, LinkedIn Learning. (2024)",
    "Learn Stylized Game Art Creation: Blender and Painter, Udemy. (2023)",
    "Unreal Engine 5 \u2013 Make AAA Game Vehicles, Udemy. (2023)",
    "100 Days of Code: The Complete Python Pro Bootcamp, Udemy. (2023)",
    "Unreal Engine 4 C++ \u2013 The Ultimate Game Developer Course, Udemy. (2023)",
    "AI System in Unreal Engine 5 and C++, Beginner to Advanced, Udemy. (2023)",
    "C++ Best Practices for Developers, LinkedIn Learning. (2022)",
    "C++ Templates and the STL, LinkedIn Learning. (2022)",
    "Learning C++, LinkedIn Learning. (2022)",
    "Master AR/VR with Unreal Engine, LinkedIn Learning. (2022)",
    "Unreal: AR Visualization 03 \u2013 Complex Interactivity, LinkedIn Learning. (2022)",
    "Unreal: AR Visualization 02 \u2013 Basic Interactivity, LinkedIn Learning. (2022)",
    "Unreal: AR Visualization 01 \u2013 Basic Concepts, LinkedIn Learning. (2022)",
    "Unreal Engine: ArchViz Design Techniques, LinkedIn Learning. (2022)",
    "Unreal Essential Training, LinkedIn Learning. (2022)",
    "Virtual Reality Foundations, LinkedIn Learning. (2022)",
    "Hard Surface Modelling in Maya 2022, Udemy. (2022)",
    "Maya 3D Masterclass \u2013 Modeling a 3D Sci-Fi Vehicle in Maya, Udemy. (2022)",
    "Unreal VR Dev: Make VR Experiences with Unreal Engine in C++, Udemy. (2022)",
    "Hard Surface Vehicle 3D Modeling in Maya, Udemy. (2022)",
    "Unreal Engine C++ Developer: Learn C++ and Make Video Games, Udemy. (2022)",
    "Unreal Engine VR Development Fundamentals, Udemy. (2022)",
    "Web Design Responsivo sem Framework, Udemy. (2021)",
]
set_all(paras[135], COURSES[0])
set_all(paras[136], COURSES[1])
_cref = paras[136]
for _c in COURSES[2:]:
    _cref = clone_after(_cref, _c)


# =====================================================================
# 13.  LANGUAGES
# =====================================================================
set_all(paras[143], "\tEnglish: Fluent (C2)")
set_all(paras[144], "Portuguese: Native")
clone_after(paras[144], "Swedish: Beginner (willing to learn)")


# =====================================================================
# 15.  OTHER INFORMATION — Interests
# =====================================================================
clear_all(paras[160])  # remove optional Family line
set_all(paras[164],
        "Motivated by: building tools that remove friction for other engineers, "
        "and shipping real products end to end.")
set_all(paras[165],
        "Activities: game & tools development, contributing plugins to the "
        "Unreal/FAB community, and continuous learning.")
set_all(paras[166],
        "Appreciate: clean, well-measured engineering, cross-disciplinary teams, "
        "and work with tangible impact.")


# =====================================================================
# FORMATTING CLEANUP — remove oversized gaps while keeping the design
# =====================================================================
# 1) The template has a next-page section break right after the Ford occlusion
#    bullet, which pushes the rest of the CV onto a new page (the big gap in the
#    print). Convert every in-body section break to a continuous break so content
#    flows without page jumps (headers/footers/section structure are preserved).
for _p in doc.paragraphs:
    _pPr = _p._p.find(qn('w:pPr'))
    if _pPr is None:
        continue
    _sect = _pPr.find(qn('w:sectPr'))
    if _sect is None:
        continue
    _t = _sect.find(qn('w:type'))
    if _t is None:
        _t = _sect.makeelement(qn('w:type'), {})
        _pgSz = _sect.find(qn('w:pgSz'))
        if _pgSz is not None:
            _pgSz.addprevious(_t)
        else:
            _sect.append(_t)
    _t.set(qn('w:val'), 'continuous')

# 2) Drop manual page breaks in the body (e.g. the one before Languages).
for _p in doc.paragraphs:
    for _br in list(_p._p.iter(qn('w:br'))):
        if _br.get(qn('w:type')) == 'page':
            _br.getparent().remove(_br)


# 3) Collapse any run of 3+ consecutive empty paragraphs down to 2 blank lines.
def _is_blank(p):
    if p.text.strip() != "":
        return False
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
        return False  # never drop a paragraph that carries a section break
    return True


_body_ps = doc.paragraphs
_run_start = None
_to_remove = []
for _idx, _p in enumerate(_body_ps + [None]):
    if _p is not None and _is_blank(_p):
        if _run_start is None:
            _run_start = _idx
    else:
        if _run_start is not None:
            if _idx - _run_start > 2:
                for _k in range(_run_start + 2, _idx):
                    _to_remove.append(_body_ps[_k])
            _run_start = None
for _p in _to_remove:
    _p._p.getparent().remove(_p._p)


doc.save(OUTPUT)
print("Saved:", OUTPUT)
