"""
Tailors the resume DOCX for a UI Programmer position (Epic Games - Fortnite),
preserving all original styling/formatting.

Target role: UI Programmer - C++, UI development on PC/console/mobile, UMG &
Slate, gameplay/UI intersections, cross-platform UI compatibility, collaboration
with art/design/product, maintainable & documented code.

Honesty note: emphasis on SkateNation XL, where the candidate owned a large
share of the game's UI screens and their interaction mechanics, including
ensuring cross-platform compatibility across PC/console. Real shipped C++/Unreal
work, including hands-on UMG & Slate UI development (multiple full menu flows).
No fabricated Fortnite or Epic experience.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer-UI-Programmer.docx"

doc = Document(INPUT)

# === SUBTITLE LINE ===
subtitle_new = "UI Programmer | C++ / Unreal Engine (UMG & Slate) | EU Citizen | +55(75)999823805 | igonildo7@gmail.com "

# === PARAGRAPH REPLACEMENTS (full-paragraph exact match) ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "UI Programmer / Software Engineer with 5 years of experience building real-time, player-facing systems in C++ and Unreal Engine. On SkateNation XL (shipped to PlayStation 5 and Xbox Series X|S), I owned a large share of the game's UI screens and their interaction mechanics, including ensuring consistent behavior and compatibility across multiple platforms. Strong command of C++ and UI/UX implementation, with hands-on Unreal UI development (UMG and Slate) building complete menu flows, gameplay-UI integration, and performance optimization. Comfortable collaborating closely with art, design, and product in remote, cross-disciplinary teams, and writing reliable, documented, maintainable code. EU Citizen open to relocation or remote work.",

    # Job titles
    "Software engineer/C++ developer, Ford Motor Company":
        "Software Engineer / C++ Developer, Ford Motor Company",
    "Software engineer, Cafundo Creative Studio":
        "Software Engineer (UI / Real-Time Interaction), Cafundo Creative Studio",
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "UI / Gameplay Programmer (C++ / Unreal Engine), Blue Gravity Studios",

    # Ford bullets
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance C++ tooling and interactive interfaces, applying strong technical fundamentals and clean, maintainable code, contributing to an estimated ~40% annual material cost reduction.",
    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Collaborated closely with internal customers (Design and Engineering teams) to design and implement the interfaces and tools they needed, reducing iteration cycles by ~50%.",
    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Optimized complex UI/visualization rendering and performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).",
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust features with strong attention to detail on functionality, readability, performance, and long-term maintainability.",

    # Cafundo bullets
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected an AI-powered interactive UI application for high-traffic public environments, serving over 500 daily interactions with zero downtime.",
    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Integrated backend services into the UI layer and optimized the real-time interaction loop, reducing response latency by 30% for a smooth, responsive user experience.",
    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Designed intuitive UI/UX systems grounded in UI/UX design principles, focused on accessibility and seamless real-time user interaction.",

    # Blue Gravity bullets
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Built and owned numerous complete menu flows, HUD, and gameplay-facing UI screens and their interaction mechanics on SkateNation XL using C++/Unreal Engine (UMG & Slate), integrating them cleanly with core gameplay features.",
    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Developed reliable, extensible UI/gameplay systems with documentation and clean architecture, improving code modularity and reducing bug-fixing overhead by 20% for the team.",
    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Ensured UI and input compatibility across multiple platforms (PC, PlayStation 5, Xbox Series X|S), adapting layouts and interaction mechanics for different control schemes and reducing input latency by 15ms.",
    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Collaborated with artists, designers, and programmers to refine UI feel and optimize performance, enhancing responsiveness on lower-end and constrained hardware.",
}

# Substring replacements (applied within a paragraph, keeps surrounding text)
substring_replacements = {
    "C# / .NET 9 / React / TypeScript / PostgreSQL":
        "C++ / Unreal Engine (UMG & Slate) / Cross-Platform UI / Gameplay-UI Integration",
}


def set_paragraph_text(para, new_text):
    """Replace full paragraph text, keeping the first run's formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def apply_substring(para, old, new):
    """Replace a substring inside the run that contains it."""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if old in para.text and para.runs:
        combined = para.text.replace(old, new)
        para.runs[0].text = combined
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


# Apply paragraph + subtitle replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Subtitle line (has email + LINKS)
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs:
            para.runs[0].text = subtitle_new
            if len(para.runs) >= 5:
                para.runs[4].text = "\n"
            if len(para.runs) >= 6:
                para.runs[5].text = "\n"
            if len(para.runs) >= 7:
                para.runs[6].text = "LINKS"
            for i, run in enumerate(para.runs[1:], start=1):
                if i not in (4, 5, 6):
                    run.text = ""
        for hl in para._element.findall(qn('w:hyperlink')):
            for r in hl.findall(qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None:
                    t.text = ""
        continue

    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            set_paragraph_text(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in substring_replacements.items():
        if old in para.text:
            apply_substring(para, old, new)
            break


# === SKILLS TABLE ===
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()

        if "Languages:" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Languages:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "C++ (modern, performance-critical), C#, Python, HLSL, SQL")

        elif "Systems" in label:
            set_paragraph_text(cells[0].paragraphs[0], "UI & Engine:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "Unreal Engine 5 (C++/Blueprint), UMG & Slate, complete menu flows, HUD, "
                "gameplay-UI integration, cross-platform UI (PC/console/mobile), "
                "UI/UX design principles, input & interaction systems")

        elif "Cloud" in label or "DevOps" in label:
            labels = ["Performance & Quality:", "Tools & Workflow:", "Engineering practices:"]
            values = [
                "CPU/GPU profiling & UI optimization, frame-rate & memory budgeting, "
                "bug fixing, attention to detail on functionality, readability & compliance",
                "Git/Perforce concepts, Visual Studio, Figma, Blender, CI/CD, AI-assisted development",
                "Reliable, maintainable & documented code, extensible design, clean architecture, "
                "cross-disciplinary collaboration (art/design/product), remote teamwork, C2 English",
            ]
            for pi, para in enumerate(cells[0].paragraphs):
                set_paragraph_text(para, labels[pi] if pi < len(labels) else "")
            for pi, para in enumerate(cells[1].paragraphs):
                set_paragraph_text(para, values[pi] if pi < len(values) else "")


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
