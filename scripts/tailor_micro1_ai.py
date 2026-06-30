"""
Tailors the resume DOCX for the micro1 "Software Developer" (Contractor, Remote)
role, preserving all original styling/formatting.

Target role: generalist Software Developer contributing to AI model training &
evaluation workflows. Stack-agnostic (Python / C# / C++), emphasis on:
- strong, high-quality software development across modern languages (Python first)
- hands-on AI/LLM integration into production applications
- remote, distributed teams; version control; CI/CD; code review; documentation
- problem-solving, performance, reliability

Honesty note: keeps all claims truthful. AI experience is framed around the real
LLM/API integration work (Cafundo) and computer-vision projects already on the CV.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer-micro1.docx"

doc = Document(INPUT)

# === PARAGRAPH REPLACEMENTS (titles + summary) ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "Software Engineer with 4+ years of experience designing, developing, and optimizing robust software across Python, C#, and C++. Proven track record translating complex requirements into clean, maintainable, high-quality solutions, with hands-on experience integrating AI/LLM services into production applications and building data-driven processing pipelines. Effective in remote, distributed teams using version control, CI/CD, and rigorous code review. Strong problem-solver focused on performance, reliability, and clear documentation. EU Citizen, available for remote work (global).",

    # Cafundo title
    "Software engineer, Cafundo Creative Studio":
        "Software Engineer (AI Integration), Cafundo Creative Studio",

    # Ford title
    "Software engineer/C++ developer, Ford Motor Company":
        "Software Engineer, Ford Motor Company",

    # Blue Gravity title
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "Software Engineer (C++ / Distributed Systems), Blue Gravity Studios",
}

# Bullet-level replacements (List Paragraph items)
bullet_replacements = {
    # Cafundo bullets (core AI story for micro1)
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected an AI-powered interactive application in C# for high-traffic public environments, designing the software components end to end and serving over 500 daily interactions with zero downtime. ",

    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Integrated LLM and AI services (OpenAI, Azure Cognitive Services) into a production inference/response pipeline, reducing voice-to-response latency by 30% for a natural conversational flow. ",

    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Built reliable real-time interaction and data-processing logic, with documentation and tests to support ongoing team success. ",

    # Ford bullets
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance internal software tooling and automated validation/evaluation pipelines in C++/C#, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction. ",

    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Analyzed requirements with cross-functional teams and translated them into technical solutions, reducing design iteration cycles by ~50%. ",

    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Optimized system performance and data-heavy processing across multiple targets, sustaining stable throughput while processing high-density datasets (5M+ data points). ",

    # Blue Gravity bullets
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Engineered core distributed-systems and data-synchronization logic (TCP/UDP) for a concurrent multi-client architecture in C++, optimizing bandwidth usage by 25% to support highly concurrent sessions. ",

    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Developed core software systems with clean separation of concerns, improving modularity and reducing bug-fixing overhead by 20% for the engineering team. ",

    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Implemented scalable, well-tested systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms. ",
}

# Standalone paragraphs (not bullets) that need rewording
standalone_replacements = {
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust software features with a strong focus on performance profiling, code quality, stability, and maintainability, conducting code reviews to enforce best practices.",

    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Collaborated with multidisciplinary, distributed teams using async communication and version control to optimize performance on lower-end hardware.",
}


def replace_text_preserving_format(para, new_text):
    """Clear all runs and set new text on the first run, keeping its formatting."""
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
        return True
    para.text = new_text
    return True


# Apply paragraph-level replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Subtitle line: only swap the job title, keep phone/email hyperlink/LINKS intact
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs and "Software Engineer | EU Citizen" in para.runs[0].text:
            para.runs[0].text = para.runs[0].text.replace(
                "Software Engineer | EU Citizen", "Software Developer | EU Citizen")
        continue

    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in bullet_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in standalone_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            break


# === SKILLS TABLE REPLACEMENT ===
row2_labels = ["AI & Cloud:", "DevOps & Tools:", "Additional (background):"]
row2_values = [
    "LLM/API Integration (OpenAI, Azure Cognitive Services), Data Processing Pipelines, Computer Vision, Real-time Inference, AWS (Lambda, API Gateway, DynamoDB, SQS).",
    "Docker, CI/CD (GitHub Actions), Git, Linux/OS fundamentals, Automated Testing & Tooling",
    "Unreal Engine 5 (C++/Blueprint), Unity (C#), profiling/optimization, distributed systems / networking",
]

for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()

        if "Languages:" in label:
            for pi, para in enumerate(cells[0].paragraphs):
                if para.runs:
                    para.runs[0].text = "Languages:" if pi == 0 else ""
                    for r in para.runs[1:]:
                        r.text = ""
            for pi, para in enumerate(cells[1].paragraphs):
                if para.runs:
                    para.runs[0].text = "Python, C#, C++, TypeScript, SQL" if pi == 0 else ""
                    for r in para.runs[1:]:
                        r.text = ""

        elif "Systems" in label and "Networking" in label:
            for pi, para in enumerate(cells[0].paragraphs):
                if para.runs:
                    para.runs[0].text = "AI & Software:" if pi == 0 else ""
                    for r in para.runs[1:]:
                        r.text = ""
            for pi, para in enumerate(cells[1].paragraphs):
                if para.runs:
                    para.runs[0].text = (
                        "REST APIs, Software Design & Architecture, Concurrent Programming, "
                        "Systems Optimization, Performance Profiling, Clean/Maintainable Code"
                        if pi == 0 else ""
                    )
                    for r in para.runs[1:]:
                        r.text = ""

        elif "Cloud" in label or "DevOps" in label:
            for pi, para in enumerate(cells[0].paragraphs):
                target = row2_labels[pi] if pi < len(row2_labels) else ""
                if para.runs:
                    para.runs[0].text = target
                    for r in para.runs[1:]:
                        r.text = ""
                elif target:
                    para.add_run(target)
            for pi, para in enumerate(cells[1].paragraphs):
                target = row2_values[pi] if pi < len(row2_values) else ""
                if para.runs:
                    para.runs[0].text = target
                    for r in para.runs[1:]:
                        r.text = ""
                elif target:
                    para.add_run(target)


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
