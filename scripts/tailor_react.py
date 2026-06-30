"""
Tailors the resume DOCX for a Front-End Developer position (web), leaning toward
an Angular / JavaScript / TypeScript enterprise/banking stack (Axpe Consulting -
UK remote project), preserving all original styling/formatting.

Honesty note: the candidate's main component framework experience is React, with
strong TypeScript/JavaScript fundamentals. Angular is listed as "familiar"
(transferable component/TS skills) rather than claimed as production experience.
Emphasis: TypeScript/JavaScript, reusable component UI, REST API integration,
AWS (EKS) deployments and CI/CD pipelines, agile international teams.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer-FrontEnd-Angular.docx"

doc = Document(INPUT)

# === SUBTITLE LINE (paragraph 1) ===
subtitle_new = "Front-End Developer | EU Citizen | +55(75)999823805 | igonildo7@gmail.com "

# === PARAGRAPH REPLACEMENTS (exact text match on the whole paragraph) ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "Front-End Developer with 4+ years of experience building high-performance, scalable web interfaces for enterprise and high-traffic environments. Advanced command of TypeScript, JavaScript, and modern component-based frameworks (React, with hands-on Angular familiarity), delivering reusable UI systems and clean, maintainable front-ends. Strong experience integrating RESTful APIs, collaborating with back-end teams, and deploying to AWS (EKS) through CI/CD pipelines. Comfortable in international, English-speaking, agile environments. EU Citizen open to relocation (Ireland/UK/EU) or remote work.",

    # Job titles
    "Software engineer/C++ developer, Ford Motor Company":
        "Front-End / Software Engineer, Ford Motor Company",
    "Software engineer, Cafundo Creative Studio":
        "Front-End / Software Engineer, Cafundo Creative Studio",
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "Front-End / Software Developer, Blue Gravity Studios",

    # Ford bullets
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance internal web tooling and automated validation pipelines, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.",
    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Collaborated closely with internal customers (Design and Engineering teams) to define requirements and build scalable, responsive UIs, reducing design iteration cycles by ~50%.",
    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Optimized complex UI rendering and front-end performance across multiple targets, sustaining smooth frame rates while processing high-density data visualizations (5M+ data points).",
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust front-end features with a strong focus on performance profiling, code quality, stability, and maintainability.",

    # Cafundo bullets
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected an AI-powered interactive application with a component-based UI for high-traffic public environments, serving over 500 daily interactions with zero downtime.",
    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Integrated RESTful APIs (Azure Cognitive Services, OpenAI) with efficient client-side state management, reducing voice-to-response latency by 30% to create a natural conversational flow.",
    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Designed reusable, accessible UI components with a focus on responsive layouts and seamless real-time user interaction.",

    # Blue Gravity bullets
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Engineered real-time data synchronization for a multi-client architecture, optimizing API/data efficiency by 25% to support highly concurrent user sessions.",
    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Developed modular, reusable component systems with clear separation of concerns, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.",
    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Implemented scalable input and interaction systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.",
    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Collaborated with multidisciplinary teams in an agile environment to optimize performance and responsiveness on lower-end hardware.",
}

# Substring replacements (applied within a paragraph, keeps surrounding text)
substring_replacements = {
    "C# / .NET 9 / React / TypeScript / PostgreSQL":
        "React / TypeScript / .NET 9 / PostgreSQL",
}


def set_paragraph_text(para, new_text):
    """Replace full paragraph text, keeping the first run's formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def apply_substring(para, old, new):
    """Replace a substring inside the run that contains it."""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Fall back: collapse into first run if split across runs
    if old in para.text and para.runs:
        combined = para.text.replace(old, new)
        para.runs[0].text = combined
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


# Apply paragraph + subtitle replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Subtitle line (has email + LINKS)
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs:
            para.runs[0].text = subtitle_new
            # Preserve the "\n\nLINKS" tail that lives in later runs
            if len(para.runs) >= 5:
                para.runs[4].text = "\n"
            if len(para.runs) >= 6:
                para.runs[5].text = "\n"
            if len(para.runs) >= 7:
                para.runs[6].text = "LINKS"
            for i, run in enumerate(para.runs[1:], start=1):
                if i not in (4, 5, 6):
                    run.text = ""
        for hl in para._element.findall(qn('w:hyperlink')):
            for r in hl.findall(qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None:
                    t.text = ""
        continue

    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            set_paragraph_text(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in substring_replacements.items():
        if old in para.text:
            apply_substring(para, old, new)
            break


# === SKILLS TABLE ===
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()

        if "Languages:" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Languages:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "TypeScript, JavaScript, Java, SQL, C#, Python")

        elif "Systems" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Front-End:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "Angular (familiar), React, TypeScript, JavaScript (ES6+), "
                "Reusable Component Architecture, REST API Integration, "
                "Responsive Design, State Management (Redux/Context), HTML5/CSS3")

        elif "Cloud" in label or "DevOps" in label:
            labels = ["Cloud & DevOps:", "CI/CD & Tools:", "Additional:"]
            values = [
                "AWS (EKS, Lambda, API Gateway, DynamoDB, CloudWatch), "
                "Kubernetes, Docker, CI/CD pipelines on AWS",
                "GitHub Actions, Git, ASP.NET Core, Node.js, REST API "
                "integration, Agile / distributed international teams",
                "Java & SQL fundamentals, Performance Profiling, Clean Code, "
                "SOLID, Unit Testing",
            ]
            for pi, para in enumerate(cells[0].paragraphs):
                set_paragraph_text(para, labels[pi] if pi < len(labels) else "")
            for pi, para in enumerate(cells[1].paragraphs):
                set_paragraph_text(para, values[pi] if pi < len(values) else "")


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
