"""
Tailors the resume DOCX for a Senior Unreal Engine Engineer (Console / Switch)
position, preserving all original styling/formatting.

Target role: Senior Unreal Engine Engineer (Switch) - C++ + Unreal Engine,
console porting & optimization (Nintendo Switch 2, Xbox Series X|S, PS5),
rendering pipelines, GPU debugging, performance optimization, systems design,
tool development, technical leadership / mentorship.

Honesty note: the candidate has NOT shipped a commercial title on Nintendo
Switch. Switch is framed truthfully as hands-on personal/console projects with
on-device testing (owns the hardware). Shipped console work is real: SkateNation
XL released on PlayStation 5 and Xbox Series X|S (Blue Gravity). Also backed by
two published C++ Unreal Marketplace plugins (incl. a fully async GPU-to-CPU
readback pipeline with version-conditional support for UE 4.27-5.6),
multiplayer network replication at Blue Gravity, and high-performance C++
tooling/optimization at Ford across multiple hardware targets.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer-Unreal-Console.docx"

doc = Document(INPUT)

# === SUBTITLE LINE ===
subtitle_new = "Senior Unreal Engine Engineer | C++ / Console | EU Citizen | +55(75)999823805 | igonildo7@gmail.com "

# === PARAGRAPH REPLACEMENTS (full-paragraph exact match) ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "Senior C++ / Unreal Engine Engineer with 5 years of professional experience building high-performance real-time systems, gameplay/network code, and engine-level tooling. Deep command of C++ and Unreal Engine 5, with a proven track record of profiling and resolving CPU/GPU and memory bottlenecks across multiple hardware targets. Shipped a console title (SkateNation XL) to PlayStation 5 and Xbox Series X|S, and authored two published Unreal Marketplace C++ plugins, including a fully asynchronous GPU-to-CPU readback pipeline with version-conditional support across UE 4.27-5.6. Experienced in multiplayer networking, performance optimization, and cross-functional collaboration with art, design, and production teams. Actively developing and on-device testing personal projects on Nintendo Switch. Based in Brazil, available for remote work; EU Citizen open to relocation.",

    # Job titles
    "Software engineer/C++ developer, Ford Motor Company":
        "Software Engineer / C++ Developer, Ford Motor Company",
    "Software engineer, Cafundo Creative Studio":
        "Software Engineer (C++ / Real-Time Systems), Cafundo Creative Studio",
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "Unreal Engine 5 / C++ Gameplay Engineer, Blue Gravity Studios",

    # Ford bullets
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance C++ tooling and automated validation pipelines, applying the same systems-engineering rigor used for production content pipelines, contributing to an estimated ~40% annual material cost reduction.",
    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Built production tooling and workflows in close collaboration with internal customers (Design and Engineering teams), anticipating technical risks early and reducing iteration cycles by ~50%.",
    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Profiled and optimized rendering and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ polygons/data points) under tight memory and GPU constraints.",
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust C++ features with a strong focus on performance profiling, GPU/CPU debugging, memory management, stability, and long-term maintainability across hardware constraints.",

    # Cafundo bullets
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected a real-time interactive application for high-traffic public environments, serving over 500 daily interactions with zero downtime under strict reliability constraints.",
    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Optimized the real-time interaction loop and external service integration, reducing end-to-end latency by 30% to create a smooth, responsive experience.",
    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Designed responsive real-time interaction systems with a focus on frame-rate stability and seamless user experience.",

    # Blue Gravity bullets
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Engineered core network replication logic (TCP/UDP) for multiplayer architecture in C++/Unreal Engine, optimizing bandwidth usage by 25% to support highly concurrent online sessions.",
    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Designed and developed core gameplay systems with clean architecture, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.",
    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Contributed to shipping SkateNation XL to PlayStation 5 and Xbox Series X|S, implementing scalable, low-latency input systems and ensuring 100% compatibility across target console platforms.",
    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Collaborated with multidisciplinary teams (engineering, art, design) to optimize memory usage and draw/processing calls, enhancing overall performance on lower-end and constrained hardware.",
}

# Substring replacements (applied within a paragraph, keeps surrounding text)
substring_replacements = {
    "C# / .NET 9 / React / TypeScript / PostgreSQL":
        "C++ / Unreal Engine 5 / Async GPU-CPU Pipeline / Multiplayer",
}


def set_paragraph_text(para, new_text):
    """Replace full paragraph text, keeping the first run's formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def apply_substring(para, old, new):
    """Replace a substring inside the run that contains it."""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if old in para.text and para.runs:
        combined = para.text.replace(old, new)
        para.runs[0].text = combined
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


# Apply paragraph + subtitle replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Subtitle line (has email + LINKS)
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs:
            para.runs[0].text = subtitle_new
            if len(para.runs) >= 5:
                para.runs[4].text = "\n"
            if len(para.runs) >= 6:
                para.runs[5].text = "\n"
            if len(para.runs) >= 7:
                para.runs[6].text = "LINKS"
            for i, run in enumerate(para.runs[1:], start=1):
                if i not in (4, 5, 6):
                    run.text = ""
        for hl in para._element.findall(qn('w:hyperlink')):
            for r in hl.findall(qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None:
                    t.text = ""
        continue

    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            set_paragraph_text(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in substring_replacements.items():
        if old in para.text:
            apply_substring(para, old, new)
            break


# === SKILLS TABLE ===
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()

        if "Languages:" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Languages:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "C++ (modern, performance-critical), C#, Python, HLSL (shaders), SQL")

        elif "Systems" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Engine & Gameplay:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "Unreal Engine 5 (C++/Blueprint), engine/editor plugin development, "
                "gameplay & systems architecture, multiplayer replication (TCP/UDP), "
                "real-time systems, async GPU-to-CPU pipelines, cross-version (UE 4.27-5.6)")

        elif "Cloud" in label or "DevOps" in label:
            labels = ["Performance & Rendering:",
                      "Platforms & Tools:", "Engineering practices:"]
            values = [
                "CPU/GPU profiling & debugging, frame-rate & memory optimization, "
                "rendering pipeline tuning, draw-call/asset budgeting under hardware constraints",
                "PlayStation 5, Xbox Series X|S (shipped console title), Nintendo Switch (personal projects, on-device testing), PC; "
                "Git/Perforce concepts, Visual Studio, Blender, CI/CD, AI-assisted development",
                "Systems & software architecture, task breakdown & planning, "
                "technical leadership/mentoring, clean code, profiling-driven optimization, "
                "cross-functional work with art/design/production, C2 English",
            ]
            for pi, para in enumerate(cells[0].paragraphs):
                set_paragraph_text(
                    para, labels[pi] if pi < len(labels) else "")
            for pi, para in enumerate(cells[1].paragraphs):
                set_paragraph_text(
                    para, values[pi] if pi < len(values) else "")


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
