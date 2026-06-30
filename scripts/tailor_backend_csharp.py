"""
Tailors the resume DOCX for a Backend Software Engineer (C# / .NET) position,
preserving all original styling/formatting.

Target role: Backend Software Engineer C# (Techifide) - C#/.NET, PostgreSQL,
distributed systems / message queues (SQS), AWS, observability, CI/CD,
with graphics/3D (CAD, Three.js/WebGL) and real-time collaboration as a plus.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\IgorGarcia_Software_Development_Engineer-Backend-CSharp.docx"

doc = Document(INPUT)

# === PARAGRAPH REPLACEMENTS ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "Backend Software Engineer with 4+ years of experience designing and building scalable services in C# / .NET, distributed systems, and high-performance tooling. Proven track record shipping production-grade REST APIs, integrating external services with resilient auth/retry/rate-limit handling, and optimizing data-heavy workloads on PostgreSQL. Strong background in AWS (Lambda, SQS, API Gateway), asynchronous messaging, and observability, complemented by deep graphics/3D engineering experience (Unreal, Unity, WebGL/CAD-style tooling, computer vision). Based in Brazil, available for remote work (LATAM/global); EU Citizen open to relocation.",

    # Cafundo title
    "Software engineer, Cafundo Creative Studio":
        "Software Engineer (C# / Backend Integration), Cafundo Creative Studio",

    # Ford title
    "Software engineer/C++ developer, Ford Motor Company":
        "Backend / Software Engineer, Ford Motor Company",

    # Blue Gravity title
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "Software Engineer (Distributed Systems / C++), Blue Gravity Studios",
}

# Bullet-level replacements (List Paragraph items)
bullet_replacements = {
    # Cafundo bullets
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected an AI-powered interactive application in C# for high-traffic public environments, serving over 500 daily interactions with zero downtime.",

    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Integrated external REST APIs (Azure Cognitive Services, OpenAI) with resilient handling of auth flows, rate limits, and failure scenarios, reducing voice-to-response latency by 30%.",

    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Designed real-time, event-driven interaction systems focused on reliability and a seamless user experience.",

    # Ford bullets
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance internal services and automated validation pipelines in C++/C#, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.",

    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Worked cross-functionally with internal customers (Design and Engineering teams) to define APIs and build the tooling they needed, reducing design iteration cycles by ~50%.",

    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Optimized system performance and data-heavy processing across multiple targets using profiling and monitoring, sustaining stable throughput on high-density datasets (5M+ data points).",

    # Blue Gravity bullets
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Engineered core distributed-systems and network synchronization logic (TCP/UDP) for concurrent multi-client architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent sessions.",

    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Developed core backend systems following clean OOP design, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.",

    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Implemented scalable, low-latency processing systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.",
}

# Standalone paragraph replacements
standalone_replacements = {
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust backend features with a strong focus on performance profiling, observability, stability, and long-term maintainability.",

    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Collaborated with multidisciplinary teams in an agile environment to optimize resource usage and processing throughput, enhancing overall system performance.",
}


def replace_text_preserving_format(para, new_text):
    """Clear all runs and set new text on the first run, preserving its format."""
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
        return True
    return False


# Apply paragraph replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Special handling for subtitle line
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = "Backend Software Engineer | C# / .NET | EU Citizen | +55(75)999823805 | igonildo7@gmail.com "
            if len(para.runs) >= 5:
                para.runs[4].text = "\n"
            if len(para.runs) >= 6:
                para.runs[5].text = "\n"
            if len(para.runs) >= 7:
                para.runs[6].text = "LINKS"
        for hl in para._element.findall(qn('w:hyperlink')):
            for r in hl.findall(qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None:
                    t.text = ""
        continue

    # Direct paragraph replacements (summary + titles)
    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            matched = True
            break
    if matched:
        continue

    # Bullet replacements
    for old, new in bullet_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            matched = True
            break
    if matched:
        continue

    # Standalone replacements
    for old, new in standalone_replacements.items():
        if old.strip() == text:
            replace_text_preserving_format(para, new)
            break


# === TABLE (SKILLS) REPLACEMENT ===
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            label = cells[0].text.strip()

            if "Languages:" in label and "C++" in cells[1].text:
                for pi, para in enumerate(cells[1].paragraphs):
                    if para.runs:
                        para.runs[0].text = "C#, .NET, SQL, Python, C++, TypeScript" if pi == 0 else ""
                        for r in para.runs[1:]:
                            r.text = ""

            elif "Systems" in label and "Networking" in label:
                for pi, para in enumerate(cells[0].paragraphs):
                    if para.runs:
                        para.runs[0].text = "Backend & APIs:" if pi == 0 else ""
                        for r in para.runs[1:]:
                            r.text = ""
                for pi, para in enumerate(cells[1].paragraphs):
                    if para.runs:
                        para.runs[0].text = (
                            "ASP.NET Core, REST & GraphQL APIs, Entity Framework Core, "
                            "PostgreSQL (schema design, indexing, query optimization), "
                            "Microservices, Distributed Systems, Concurrent/Async Programming"
                        ) if pi == 0 else ""
                        for r in para.runs[1:]:
                            r.text = ""

            elif "Cloud" in label or "DevOps" in label:
                labels = ["Cloud & Messaging:", "DevOps & Tools:", "Graphics & 3D (background):"]
                values = [
                    "AWS (Lambda, API Gateway, SQS, DynamoDB, CloudWatch), Message Queues (SQS / Kafka concepts), Serverless, Observability (logging / tracing / monitoring)",
                    "Docker, CI/CD (GitHub Actions / GitLab), Git, Linux fundamentals, AI-assisted development",
                    "Unreal Engine 5 (C++/Blueprint), Unity (C#), WebGL / 3D & CAD-style tooling, computer vision, real-time systems",
                ]
                for pi, para in enumerate(cells[0].paragraphs):
                    if para.runs:
                        para.runs[0].text = labels[pi] if pi < len(labels) else ""
                        for r in para.runs[1:]:
                            r.text = ""
                    elif pi < len(labels):
                        para.add_run(labels[pi])
                for pi, para in enumerate(cells[1].paragraphs):
                    if para.runs:
                        para.runs[0].text = values[pi] if pi < len(values) else ""
                        for r in para.runs[1:]:
                            r.text = ""
                    elif pi < len(values):
                        para.add_run(values[pi])


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
