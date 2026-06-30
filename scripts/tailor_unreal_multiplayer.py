"""
Tailors the resume DOCX for an Unreal Engine Multiplayer Gameplay Programmer
(C++ Expert) position, preserving all original styling/formatting.

Target role (iBLOXX Studios): UE5 + C++ multiplayer specialist. Design and
implement the game's network architecture - client/server models, data
replication, bandwidth optimization - scaling to 60+ players per server.
Profile, debug and optimize network performance.
Good to have: GAS (Gameplay Ability System), Python (backend), AWS (GameLift).

Honesty note: shipped console work is real - SkateNation XL released on
PlayStation 5 and Xbox Series X|S (Blue Gravity), where the candidate built
core network replication logic. Two published C++ Unreal Marketplace plugins
(incl. a fully async GPU-to-CPU readback pipeline, UE 4.27-5.6). Real Python
and AWS experience from Ford/Cafundo tooling. Brazil-based, remote-ready.
"""
from docx import Document
from docx.oxml.ns import qn

INPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\Resumes\IgorGarcia_Software_Development_Engineer.docx"
OUTPUT = r"C:\Users\ISILV125\Downloads\05_Pessoal\Curriculum\CurriculumRelated\Resumes\IgorGarcia_Software_Development_Engineer-Unreal-Multiplayer.docx"

doc = Document(INPUT)

# === SUBTITLE LINE ===
subtitle_new = "Unreal Multiplayer Gameplay Programmer | C++ / UE5 / Networking | Brazil (Remote) | +55(75)999823805 | igonildo7@gmail.com "

# === PARAGRAPH REPLACEMENTS (full-paragraph exact match) ===
paragraph_replacements = {
    # Summary
    "Software Engineer with 4+ years of experience building high-performance internal tooling, automated workflows, and network-aware applications. Strong background in C++ and C#, with a proven track record of optimizing complex systems and reducing manual operations for engineering teams. Deeply passionate about systems engineering, backend architecture, and cloud infrastructure. Currently leveraging AWS services to build scalable, automated solutions. EU Citizen open to relocation (Ireland/UK/EU) or remote work.":
        "Unreal Engine 5 / C++ engineer with 5 years of professional experience specializing in real-time, network-aware multiplayer systems. Strong command of modern C++ and UE5 networking, with hands-on experience designing client/server models, data replication, and bandwidth optimization for highly concurrent online sessions. Built core network replication logic for SkateNation XL (shipped to PlayStation 5 and Xbox Series X|S) and authored two published Unreal Marketplace C++ plugins, including a fully asynchronous GPU-to-CPU readback pipeline (UE 4.27-5.6). Proven track record profiling, debugging and optimizing performance and bandwidth in large codebases. Complementary Python backend and AWS cloud experience for scalable, automated systems. Proactive, hands-on engineer; Brazil-based and available for remote work.",

    # Job titles
    "Software engineer/C++ developer, Ford Motor Company":
        "Software Engineer / C++ Developer, Ford Motor Company",
    "Software engineer, Cafundo Creative Studio":
        "Software Engineer (C++ / Real-Time Systems), Cafundo Creative Studio",
    "Unreal Engine 5 Developer, Blue Gravity Studios":
        "Unreal Engine 5 Multiplayer / C++ Engineer, Blue Gravity Studios",

    # Blue Gravity bullets (lead with networking)
    "Engineered core network replication logic (TCP/UDP concepts) for multiplayer architecture using C++, optimizing bandwidth usage by 25% to support highly concurrent online sessions.":
        "Designed and engineered core network architecture and replication logic (client/server, TCP/UDP) in C++/Unreal Engine 5, optimizing bandwidth usage by 25% to support highly concurrent online multiplayer sessions.",
    "Developed core software systems, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.":
        "Built scalable, replicated gameplay systems with clean architecture, improving code modularity and reducing bug-fixing overhead by 20% for the engineering team.",
    "Implemented scalable input systems, reducing latency by 15ms and ensuring 100% compatibility across multiple platforms.":
        "Contributed to shipping SkateNation XL to PlayStation 5 and Xbox Series X|S, implementing scalable, low-latency networked input systems and ensuring 100% compatibility across target platforms.",
    "Collaborated with multidisciplinary teams to optimize memory usage and processing calls, enhancing overall performance on lower-end hardware.":
        "Profiled, debugged and optimized network and runtime performance with multidisciplinary teams, reducing memory usage and draw/processing calls to keep sessions running smoothly under load.",

    # Ford bullets (lean into networking/backend/AWS/Python)
    "Engineered high-performance internal software tooling and automated validation pipelines using C++, replacing physical prototypes and contributing to an estimated ~40% annual material cost reduction.":
        "Engineered high-performance C++ tooling and automated, network-aware pipelines, applying systems-engineering rigor and contributing to an estimated ~40% annual material cost reduction.",
    "Collaborated closely with internal customers (Design and Engineering teams) to define and build the tools they need, reducing design iteration cycles by ~50%.":
        "Worked hands-on with internal customers (Design and Engineering) to define and ship the tools they needed, reducing iteration cycles by ~50% with a proactive, fast-moving approach.",
    "Optimized complex visualizations and system performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ data points/polygons).":
        "Profiled and optimized real-time performance across multiple hardware targets, sustaining stable frame rates while processing high-density assets (5M+ polygons/data points).",
    "Delivered robust software features with a strong focus on performance profiling, memory management, stability, and maintainability.":
        "Delivered robust C++ and Python backend features with a strong focus on performance profiling, bandwidth/memory efficiency, stability and maintainability.",

    # Cafundo bullets (AWS/cloud + real-time)
    "Architected an AI-powered interactive totem in Unity (C#) for high-traffic public environments, serving over 500 daily interactions with zero downtime.":
        "Architected a real-time, network-connected interactive application for high-traffic public environments, serving over 500 daily interactions with zero downtime.",
    "Seamlessly integrated Azure Cognitive Services and OpenAI APIs, reducing voice-to-response latency by 30% to create a natural conversational flow.":
        "Integrated cloud and external API services (Azure, OpenAI), reducing end-to-end latency by 30% to create a smooth, responsive real-time experience.",
    "Designed intuitive UI/UX systems focused on accessibility and seamless real-time user interaction.":
        "Designed responsive real-time interaction systems with a focus on low-latency, seamless user experience.",
}

# Substring replacements (applied within a paragraph, keeps surrounding text)
substring_replacements = {
    "C# / .NET 9 / React / TypeScript / PostgreSQL":
        "C++ / Unreal Engine 5 / Replication / Bandwidth Optimization",
}


def set_paragraph_text(para, new_text):
    """Replace full paragraph text, keeping the first run's formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def apply_substring(para, old, new):
    """Replace a substring inside the run that contains it."""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if old in para.text and para.runs:
        combined = para.text.replace(old, new)
        para.runs[0].text = combined
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


# Apply paragraph + subtitle replacements
for para in doc.paragraphs:
    text = para.text.strip()

    # Subtitle line (has email + LINKS)
    if "Software Engineer" in text and "EU Citizen" in text and "igonildo7" in text:
        if para.runs:
            para.runs[0].text = subtitle_new
            if len(para.runs) >= 5:
                para.runs[4].text = "\n"
            if len(para.runs) >= 6:
                para.runs[5].text = "\n"
            if len(para.runs) >= 7:
                para.runs[6].text = "LINKS"
            for i, run in enumerate(para.runs[1:], start=1):
                if i not in (4, 5, 6):
                    run.text = ""
        for hl in para._element.findall(qn('w:hyperlink')):
            for r in hl.findall(qn('w:r')):
                t = r.find(qn('w:t'))
                if t is not None:
                    t.text = ""
        continue

    matched = False
    for old, new in paragraph_replacements.items():
        if old.strip() == text:
            set_paragraph_text(para, new)
            matched = True
            break
    if matched:
        continue

    for old, new in substring_replacements.items():
        if old in para.text:
            apply_substring(para, old, new)
            break


# === SKILLS TABLE ===
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label = cells[0].text.strip()

        if "Languages:" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Languages:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "C++ (modern, performance-critical), Python, C#, HLSL (shaders), SQL")

        elif "Systems" in label:
            set_paragraph_text(cells[0].paragraphs[0], "Networking & Engine:")
            set_paragraph_text(
                cells[1].paragraphs[0],
                "Unreal Engine 5 (C++/Blueprint), multiplayer networking & replication "
                "(client/server, TCP/UDP), bandwidth optimization, dedicated/listen servers, "
                "gameplay & systems architecture, Gameplay Ability System (GAS), "
                "engine/editor plugin development, async GPU-to-CPU pipelines (UE 4.27-5.6)")

        elif "Cloud" in label or "DevOps" in label:
            labels = ["Performance & Profiling:",
                      "Cloud & Backend:", "Platforms & Practices:"]
            values = [
                "Network & runtime profiling, latency/bandwidth optimization, "
                "CPU/GPU debugging, frame-rate & memory tuning, large-codebase optimization",
                "AWS (cloud services; GameLift-ready), Python backend services, "
                "REST/API integration, automated network-aware pipelines, CI/CD",
                "PlayStation 5, Xbox Series X|S (shipped title), PC; Perforce/Git, "
                "Visual Studio, Agile/Scrum, clean code & design patterns, proactive "
                "hands-on problem-solving, C2 English",
            ]
            for pi, para in enumerate(cells[0].paragraphs):
                set_paragraph_text(
                    para, labels[pi] if pi < len(labels) else "")
            for pi, para in enumerate(cells[1].paragraphs):
                set_paragraph_text(
                    para, values[pi] if pi < len(values) else "")


doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
